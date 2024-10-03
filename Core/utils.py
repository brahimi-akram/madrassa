from docx import Document
import os
from .models import *
def replace_placeholders_in_paragraph(paragraph, replacements):
    # Combine all runs into one single string
    full_text = ''.join(run.text for run in paragraph.runs)
    # Replace placeholders in the full text
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
    
    # Clear the paragraph runs (to replace with new text)
    for run in paragraph.runs:
        run.text = ''

    # Redistribute the modified text back into the first run
    if paragraph.runs:
        paragraph.runs[0].text = full_text


def replace_placeholders_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, replacements)

def generate_doc(student_id):
    student = Student.objects.get(id= student_id)
    current_directory = os.getcwd()
    file_path = os.path.join(current_directory,'template.docx')
    output_path = os.path.join(current_directory,'media')
    output_path = os.path.join(output_path,f'{student}.docx')
    replacement={
        'student_id':student.num,
        'teacher':student.teacher.name,
        'level':student.level.name,
        'situ':student.situation,
        'studentName':f'{student.first_name} {student.last_name}',
        'birthDate':student.date_birth,
        'address':student.adress,
        'class':f'{student.level.name} {student.level_year}',
        'parentName':student.parent.name,
        'parentPhone':student.parent.phone
    }
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacement)

    # Iterate through tables to replace placeholders
    for table in doc.tables:
        replace_placeholders_in_table(table, replacement)

    doc.save(output_path)